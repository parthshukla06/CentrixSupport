import os
import hashlib
import logging
import mimetypes
import tempfile
import zipfile
import json
import csv
import re
import nltk
import docx
import whisper
import pdfplumber
import pytesseract
import cv2
from nltk.tokenize import sent_tokenize
from langchain.schema import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_huggingface import HuggingFaceEmbeddings
from langchain.vectorstores import Chroma
from langchain_ollama import OllamaLLM

nltk.download("punkt")

# === Config ===
CHROMA_DIR = "chroma_index"
HASH_CACHE_FILE = "text_hash.txt"
EMBEDDING_MODEL = "sentence-transformers/all-MiniLM-L6-v2"
OLLAMA_MODEL = "gemma2:2b"
CHUNK_SIZE = 500
CHUNK_OVERLAP = 50
SIMILARITY_THRESHOLD = 1.4

logging.basicConfig(level=logging.INFO, format="%(asctime)s - %(levelname)s - %(message)s")

# === File Extractors ===

def extract_text_from_pdf(path):
    try:
        with pdfplumber.open(path) as pdf:
            return "\n".join([p.extract_text() or "" for p in pdf.pages])
    except Exception as e:
        logging.error(f"PDF extraction failed: {e}")
        return ""

def extract_text_from_image(path):
    try:
        img = cv2.imread(path)
        return pytesseract.image_to_string(img)
    except Exception as e:
        logging.error(f"Image OCR failed: {e}")
        return ""

def extract_text_from_docx(path):
    try:
        return "\n".join([p.text for p in docx.Document(path).paragraphs])
    except Exception as e:
        logging.error(f"DOCX extraction failed: {e}")
        return ""

def extract_text_from_txt(path):
    try:
        with open(path, "r", encoding="utf-8") as f:
            return f.read()
    except Exception as e:
        logging.error(f"TXT extraction failed: {e}")
        return ""

def extract_text_from_csv(path):
    try:
        with open(path, newline='', encoding='utf-8') as csvfile:
            return "\n".join([" | ".join(row) for row in csv.reader(csvfile)])
    except Exception as e:
        logging.error(f"CSV extraction failed: {e}")
        return ""

def extract_text_from_json(path):
    try:
        with open(path, "r", encoding="utf-8") as f:
            return json.dumps(json.load(f), indent=2)
    except Exception as e:
        logging.error(f"JSON extraction failed: {e}")
        return ""

def extract_text_from_audio(path):
    try:
        # Use Whisper directly for transcription (no pydub required).
        model = whisper.load_model("base")
        return model.transcribe(path)["text"]
    except Exception as e:
        logging.error(f"Audio transcription failed: {e}")
        return ""

def extract_text_from_zip(path):
    text = ""
    try:
        with zipfile.ZipFile(path, 'r') as zip_ref:
            with tempfile.TemporaryDirectory() as temp_dir:
                zip_ref.extractall(temp_dir)
                for root, _, files in os.walk(temp_dir):
                    for name in files:
                        full_path = os.path.join(root, name)
                        text += extract_text(full_path) + "\n"
    except Exception as e:
        logging.error(f"ZIP extraction failed: {e}")
    return text

# File Dispatcher
def extract_text(path):
    ext = os.path.splitext(path)[1].lower()
    if ext == ".pdf":
        return extract_text_from_pdf(path)
    elif ext in [".jpg", ".jpeg", ".png"]:
        return extract_text_from_image(path)
    elif ext == ".docx":
        return extract_text_from_docx(path)
    elif ext == ".txt":
        return extract_text_from_txt(path)
    elif ext == ".csv":
        return extract_text_from_csv(path)
    elif ext == ".json":
        return extract_text_from_json(path)
    elif ext in [".mp3", ".wav", ".mp4"]:
        return extract_text_from_audio(path)
    elif ext == ".zip":
        return extract_text_from_zip(path)
    else:
        logging.warning(f"Unsupported file type: {ext}")
        return ""

# === Utility Functions ===

def clean_text(text):
    return re.sub(r'\s+', ' ', text).replace("•", "-").replace("–", "-").strip()

def semantic_chunk_text(text, chunk_size=CHUNK_SIZE, overlap=CHUNK_OVERLAP):
    sentences = sent_tokenize(text)
    chunks, current, current_len = [], [], 0

    for sent in sentences:
        if current_len + len(sent) > chunk_size:
            chunks.append(" ".join(current))
            current = current[-(overlap // len(sent)):] if sent else []
            current_len = sum(len(s) for s in current)
        current.append(sent)
        current_len += len(sent)

    if current:
        chunks.append(" ".join(current))
    return chunks

def compute_hash(text):
    return hashlib.md5(text.encode("utf-8")).hexdigest()

def has_file_changed(text):
    if os.path.exists(HASH_CACHE_FILE):
        with open(HASH_CACHE_FILE, "r") as f:
            return compute_hash(text) != f.read().strip()
    return True

def cache_hash(text):
    with open(HASH_CACHE_FILE, "w") as f:
        f.write(compute_hash(text))

# === Embedding + Vector Store ===

def embed_text_and_store(text, source_name):
    splitter = RecursiveCharacterTextSplitter(chunk_size=CHUNK_SIZE, chunk_overlap=CHUNK_OVERLAP)
    docs = [Document(page_content=chunk, metadata={"source": source_name}) for chunk in splitter.split_text(text)]
    embeddings = HuggingFaceEmbeddings(model_name=EMBEDDING_MODEL)

    if os.path.exists(CHROMA_DIR):
        vs = Chroma(persist_directory=CHROMA_DIR, embedding_function=embeddings)
        vs.add_documents(docs)
    else:
        vs = Chroma.from_documents(docs, embedding=embeddings, persist_directory=CHROMA_DIR)
    cache_hash(text)
    return vs

# === RAG Main Function ===

def get_rag_context(filepaths, query):
    # Normalize input to list
    if isinstance(filepaths, str):
        filepaths = [filepaths]

    combined_texts = []
    for fp in filepaths:
        try:
            text = clean_text(extract_text(fp))
            if text:
                combined_texts.append(text)
        except Exception as e:
            logging.error(f"Failed to process {fp}: {e}")

    if not combined_texts:
        logging.warning("No valid text extracted from files.")
        return ""

    # Merge and chunk
    full_text = "\n".join(semantic_chunk_text("\n".join(combined_texts)))

    # Detect changes & embed
    if has_file_changed(full_text):
        logging.info("Changes detected, embedding and storing...")
        vectorstore = embed_text_and_store(
            full_text,
            ", ".join([os.path.basename(p) for p in filepaths])
        )
    else:
        logging.info("Using existing vector store...")
        vectorstore = Chroma(
            persist_directory=CHROMA_DIR,
            embedding_function=HuggingFaceEmbeddings(model_name=EMBEDDING_MODEL)
        )

    # Similarity search
    try:
        results = vectorstore.similarity_search_with_score(query, k=5)
        return "\n\n".join(
            [doc.page_content for doc, score in results if score <= SIMILARITY_THRESHOLD]
        )
    except Exception as e:
        logging.error(f"RAG query error: {e}")
        return ""

# === Ollama Answering Function ===

def answer_query(query, vectorstore):
    try:
        results = vectorstore.similarity_search_with_score(query, k=5)
        relevant_docs = [doc for doc, score in results if score <= SIMILARITY_THRESHOLD]
        context = "\n\n".join([doc.page_content for doc in relevant_docs]) if relevant_docs else ""

        prompt = (
            "You are an AI assistant. Use only the given context to answer.\n"
            "If answer is not in context, reply: 'Not available in the documentation.'\n\n"
            f"Context:\n{context}\n\nQuestion: {query}\nAnswer:"
        )

        ollama = OllamaLLM(model=OLLAMA_MODEL)
        response = ollama.invoke(prompt)

        source = "📚 Documents" if context.strip() else "💡 No relevant document match"
        print(f"\n{source}\n✅ Answer: {response.strip()}")

    except Exception as e:
        print(f"❌ Retrieval or Ollama error: {e}")

# === Standalone Runner ===

def main():
    file_input = input("📁 Enter file paths (comma-separated): ").strip()
    file_paths = [f.strip() for f in file_input.split(",") if os.path.exists(f.strip())]

    if not file_paths:
        print("❌ No valid file paths provided.")
        return

    raw_combined_text = "\n".join([clean_text(extract_text(p)) for p in file_paths])
    chunks = semantic_chunk_text(raw_combined_text)
    full_text = "\n".join(chunks)

    if has_file_changed(full_text):
        vectorstore = embed_text_and_store(full_text, ", ".join([os.path.basename(p) for p in file_paths]))
    else:
        vectorstore = Chroma(persist_directory=CHROMA_DIR, embedding_function=HuggingFaceEmbeddings(model_name=EMBEDDING_MODEL))

    print("\n💬 Ask your questions (type 'exit' to quit):")
    while True:
        question = input("Question: ").strip()
        if question.lower() in ["exit", "quit"]:
            break
        answer_query(question, vectorstore)

if __name__ == "__main__":
    main()
